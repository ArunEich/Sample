import pandas as pd
import os
import asyncio
from datetime import datetime
from io import BytesIO
from concurrent.futures import ThreadPoolExecutor
from docx import Document as DocxDocument
from docx.shared import RGBColor
from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.core.credentials import AzureKeyCredential
from openai import AzureOpenAI
import fitz  # PyMuPDF
from PIL import Image
import base64
import tempfile
import platform
import subprocess

# ==================== CONFIG ====================
# ===== Azure Document Intelligence =====
AZURE_ENDPOINT = ""
AZURE_KEY = ""
document_client = DocumentAnalysisClient(endpoint=AZURE_ENDPOINT, credential=AzureKeyCredential(AZURE_KEY))

# ===== Azure OpenAI =====
OPENAI_ENDPOINT = ""
OPENAI_KEY = ""
OPENAI_DEPLOYMENT = ""
openai_client = AzureOpenAI(
    api_key=OPENAI_KEY,
    api_version="",
    azure_endpoint=OPENAI_ENDPOINT
)

# Upload log file
UPLOAD_LOG_FILE = "upload_logs.csv"

# ==================== HELPERS ====================
def init_log_file():
    """Initialize upload log file if it doesn't exist"""
    if not os.path.exists(UPLOAD_LOG_FILE):
        df = pd.DataFrame(columns=["User", "File Name", "Upload Time"])
        df.to_csv(UPLOAD_LOG_FILE, index=False)

def log_upload(user, file_name):
    """Log file upload"""
    upload_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    new_record = pd.DataFrame([[user, file_name, upload_time]], columns=["User", "File Name", "Upload Time"])
    new_record.to_csv(UPLOAD_LOG_FILE, mode="a", header=False, index=False)

def load_upload_logs():
    """Load upload logs with upload count"""
    if os.path.exists(UPLOAD_LOG_FILE):
        df = pd.read_csv(UPLOAD_LOG_FILE)
        if not df.empty:
            df["Upload Count"] = df.groupby(["User", "File Name"]).cumcount() + 1
        else:
            df["Upload Count"] = []
        return df
    return pd.DataFrame(columns=["User", "File Name", "Upload Time", "Upload Count"])

def load_checkpoints():
    """Load checkpoints from Excel file"""
    try:
        df = pd.read_excel("Rules BD Audit Automation copy.xlsx")
        if "checkpointDesc" not in df.columns:
            raise ValueError("Excel must have 'checkpointDesc' column")
        return df["checkpointDesc"].dropna().tolist()
    except FileNotFoundError:
        # Return default checkpoints if file not found
        return [
            "Document should have BD logo",
            "Document should contain project name",
            "Document should have proper page numbering",
            "Document should have SRA information",
            "Document should follow proper formatting guidelines"
        ]

def extract_pdf_content(file_bytes: bytes):
    """Extract content from PDF using Azure Document Intelligence"""
    try:
        poller = document_client.begin_analyze_document("prebuilt-layout", file_bytes)
        result = poller.result()
        full_text = result.content if result.content else ""

        tables_text = []
        for t_index, table in enumerate(result.tables, start=1):
            tables_text.append(f"\n[Table {t_index}]\n")
            rows = {}
            for cell in table.cells:
                rows.setdefault(cell.row_index, {})[cell.column_index] = cell.content.strip().replace("\n", " ")
            for row_idx in sorted(rows.keys()):
                row = rows[row_idx]
                row_text = " | ".join([row.get(col_idx, "") for col_idx in sorted(row.keys())])
                tables_text.append(row_text)
            tables_text.append("")
        tables_combined = "\n".join(tables_text)
        return f"{full_text}\n\nExtracted Tables:\n{tables_combined}"
    except Exception as e:
        return f"Error extracting PDF content: {str(e)}"

def extract_docx_content(file_bytes: bytes):
    """Extract content from DOCX file"""
    try:
        doc = DocxDocument(BytesIO(file_bytes))
        paragraphs = []

        # Body paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                paragraphs.append(p.text.strip())

        # Headers & footers
        for section in doc.sections:
            header = section.header
            footer = section.footer
            if header:
                for p in header.paragraphs:
                    if p.text.strip():
                        paragraphs.append("[HEADER] " + p.text.strip())
                if header._element.xpath('.//w:pict | .//w:drawing'):
                    paragraphs.append("[HEADER IMAGE: Possible BD Logo]")
            if footer:
                for p in footer.paragraphs:
                    if p.text.strip():
                        paragraphs.append("[FOOTER] " + p.text.strip())
                if footer._element.xpath('.//w:pict | .//w:drawing'):
                    paragraphs.append("[FOOTER IMAGE]")

        full_text = "\n".join(paragraphs)

        # Tables
        tables_text = []
        for t_index, table in enumerate(doc.tables, start=1):
            tables_text.append(f"\n[Table {t_index}]\n")
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                tables_text.append(row_text)
            tables_text.append("")
        tables_combined = "\n".join(tables_text)
        return f"{full_text}\n\nExtracted Tables:\n{tables_combined}"
    except Exception as e:
        return f"Error extracting DOCX content: {str(e)}"

# ==================== DOCX to PDF ====================
def convert_docx_to_pdf_bytes(file_bytes):
    """Convert DOCX to PDF bytes"""
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "temp.docx")
        pdf_path = os.path.join(tmpdir, "temp.pdf")
        with open(docx_path, "wb") as f:
            f.write(file_bytes)

        try:
            if platform.system() == "Windows":
                soffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"
                subprocess.run(
                    [soffice_path, "--headless", "--convert-to", "pdf", "--outdir", tmpdir, docx_path],
                    check=True
                )
            else:
                subprocess.run(
                    ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", tmpdir, docx_path],
                    check=True
                )

            if not os.path.exists(pdf_path):
                raise FileNotFoundError("PDF not generated by LibreOffice.")

            with open(pdf_path, "rb") as f:
                return f.read()

        except Exception as e:
            print(f"DOCX → PDF conversion failed: {e}")
            return None

# ==================== PDF to Images ====================
def save_pdf_images(pdf_bytes, output_dir, file_prefix, dpi=150):
    """Save PDF pages as images"""
    os.makedirs(output_dir, exist_ok=True)
    pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
    saved_files = []

    for page_num in range(len(pdf)):
        page = pdf[page_num]
        pix = page.get_pixmap(dpi=dpi)
        file_path = os.path.join(output_dir, f"{file_prefix}_page_{page_num+1}.png")
        pix.save(file_path)
        saved_files.append(file_path)

    pdf.close()
    return saved_files

def pdf_bytes_to_images(pdf_bytes, dpi=150):
    """Convert PDF bytes to PIL images"""
    pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
    images = []
    for page_num in range(len(pdf)):
        page = pdf[page_num]
        pix = page.get_pixmap(dpi=dpi)
        img = Image.open(BytesIO(pix.tobytes("png")))
        images.append(img)
    pdf.close()
    return images

# ==================== AI Check Engines ====================
def classify_rule_type(checkpoint: str) -> str:
    """Classify rule type using AI"""
    classification_prompt = f"""
You are an assistant that classifies audit checkpoints.

Checkpoint:
"{checkpoint}"

Decide the rule type:
- text_rule: Can be verified only from text.
- visual_rule: Needs visual/layout check only (logo, page numbers, formatting, signatures).
- hybrid_rule: Needs both text + visual validation.

Answer with only one word: text_rule, visual_rule, or hybrid_rule.
"""
    try:
        response = openai_client.chat.completions.create(
            model=OPENAI_DEPLOYMENT,
            messages=[{"role": "user", "content": classification_prompt}],
            temperature=0,
            max_tokens=10
        )
        return response.choices[0].message.content.strip().lower()
    except Exception as e:
        print(f"Error classifying rule type: {e}")
        return "text_rule"  # Default fallback

def check_with_text_engine(checkpoint, document_text):
    """Check checkpoint using text content only"""
    prompt = f"""
You are auditing a document.

Document Text:
{document_text}

Rule:
{checkpoint}

Instructions:
- Check if this rule is satisfied in the text (including [HEADER]/[FOOTER]/[IMAGE] placeholders).
- End with one line:
  STATUS: PASS
  STATUS: FAIL
- If FAIL, add a short recommendation.
- no need doesn't explicitly mention logo, if you find any logo that is bd logo and SRA and project name also.
"""
    try:
        response = openai_client.chat.completions.create(
            model=OPENAI_DEPLOYMENT,
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
            max_tokens=300
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Error in text engine: {str(e)}\nSTATUS: FAIL"

def encode_image_to_base64(image):
    """Convert PIL image to base64 string"""
    buffered = BytesIO()
    image.save(buffered, format="PNG")
    return base64.b64encode(buffered.getvalue()).decode("utf-8")

def check_with_vision_engine(checkpoint, page_images):
    """Check checkpoint using visual analysis with detailed prompt + checkpoint"""
    results = []

    base_prompt = (
        "You are a document compliance checker.\n\n"
        "Checkpoint Rule:\n{checkpoint}\n\n"
        "Task: Review the provided page image and verify whether the header contains the following elements:\n"
        "1. BD Logo – should be visible in the top-left corner of the header.\n"
        "2. Document Number – a code like 'DOCxxxxxxx'.\n"
        "3. Title – the document title, e.g., 'Validation Summary Report'.\n"
        "4. Revision Number – typically shown as 'Rev', 'Revision', or 'Version', followed by a number (e.g., 1.0).\n\n"
        "Instructions:\n"
        "- Look carefully at the header region (top portion of the page).\n"
        "- For each required element, state Found or Missing and provide the actual detected value if present.\n"
        "- After listing the findings, give a STATUS: PASS if all 4 required elements are present, otherwise STATUS: FAIL.\n"
        "- Keep the format consistent.\n\n"
        
    )

    for idx, img in enumerate(page_images, start=1):
        try:
            b64_img = encode_image_to_base64(img)
            full_prompt = base_prompt.format(checkpoint=checkpoint) + f"\nCheck on page {idx}."
            
            response = openai_client.chat.completions.create(
                model=OPENAI_DEPLOYMENT,
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": full_prompt},
                        {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64_img}"}}
                    ]
                }],
                temperature=0,
                max_tokens=400
            )
            results.append(response.choices[0].message.content.strip())
        except Exception as e:
            results.append(f"Error analyzing page {idx}: {str(e)}")

    final_status = "Fail"
    for r in results:
        if "status: pass" in r.lower():
            final_status = "Pass"
            break
    return results, final_status

def check_with_hybrid_engine(checkpoint, document_text, page_images):
    """Check checkpoint using both text and visual analysis"""
    text_result = check_with_text_engine(checkpoint, document_text)
    text_status = "Pass" if "status: pass" in text_result.lower() else "Fail"
    vision_results, vision_status = check_with_vision_engine(checkpoint, page_images)
    final_status = "Pass" if (text_status == "Pass" and vision_status == "Pass") else "Fail"
    return {"text_result": text_result, "vision_results": vision_results, "status": final_status}

# ==================== MAIN CHECK LOOP ====================
async def verify_checkpoints_with_ai(checkpoints, full_text, file_bytes, file_ext):
    """Main function to verify all checkpoints"""
    results = []
    page_images = []

    # Convert to images for visual analysis
    try:
        if file_ext == "pdf":
            page_images = pdf_bytes_to_images(file_bytes, dpi=150)
        elif file_ext == "docx":
            pdf_bytes = convert_docx_to_pdf_bytes(file_bytes)
            if pdf_bytes:
                page_images = pdf_bytes_to_images(pdf_bytes, dpi=150)
    except Exception as e:
        print(f"Error converting to images: {e}")
        page_images = []

    for idx, checkpoint in enumerate(checkpoints, start=1):
        try:
            rule_type = classify_rule_type(checkpoint)
            
            if rule_type == "text_rule":
                recommendation = check_with_text_engine(checkpoint, full_text)
                status = "Pass" if "status: pass" in recommendation.lower() else "Fail"
                
            elif rule_type == "visual_rule":
                if page_images:
                    vision_results, status = check_with_vision_engine(checkpoint, page_images)
                    recommendation = "\n".join(vision_results)
                else:
                    status = "Fail"
                    recommendation = "Could not perform visual analysis - no images available"
                    
            elif rule_type == "hybrid_rule":
                if page_images:
                    hybrid_result = check_with_hybrid_engine(checkpoint, full_text, page_images)
                    recommendation = f"Text: {hybrid_result['text_result']}\nVision: {hybrid_result['vision_results']}"
                    status = hybrid_result["status"]
                else:
                    # Fallback to text-only for hybrid rules
                    recommendation = check_with_text_engine(checkpoint, full_text)
                    status = "Pass" if "status: pass" in recommendation.lower() else "Fail"
                    recommendation += "\n(Visual analysis unavailable)"
            else:
                status = "Fail"
                recommendation = f"Unrecognized rule type: {checkpoint}"

        except Exception as e:
            status = "Fail"
            recommendation = f"Error processing checkpoint: {str(e)}"

        results.append({
            "S.No": idx, 
            "Checkpoint": checkpoint, 
            "Status": status, 
            "Recommendation": recommendation
        })
    
    return results

# ==================== REPORT GENERATION ====================
def generate_report(document_name, summary, results):
    """Generate DOCX report"""
    try:
        report = DocxDocument()
        report.add_heading("Document Validation Report", 0)
        report.add_paragraph(f"Document: {document_name}")
        report.add_paragraph(f"Generated On: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        report.add_heading("Summary", level=1)
        report.add_paragraph(f"Total Checkpoints: {summary['total_checkpoints']}")
        report.add_paragraph(f"Passed: {summary['passed']}")
        report.add_paragraph(f"Failed: {summary['failed']}")
        
        report.add_heading("Detailed Results", level=1)
        table = report.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "S.No"
        hdr[1].text = "Checkpoint"
        hdr[2].text = "Status"
        hdr[3].text = "Recommendation"
        
        for r in results:
            row_cells = table.add_row().cells
            row_cells[0].text = str(r["S.No"])
            row_cells[1].text = r["Checkpoint"]
            row_cells[2].text = r["Status"]
            row_cells[3].text = r["Recommendation"]
            
            # Color coding for status
            try:
                if r["Status"] == "Pass":
                    row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                else:
                    row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
            except:
                pass  # Skip coloring if it fails
        
        file_stream = BytesIO()
        report.save(file_stream)
        file_stream.seek(0)
        return file_stream
        
    except Exception as e:
        print(f"Error generating report: {e}")
        # Return a simple report
        simple_report = DocxDocument()
        simple_report.add_heading("Document Validation Report", 0)
        simple_report.add_paragraph(f"Document: {document_name}")
        simple_report.add_paragraph(f"Generated On: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        simple_report.add_paragraph(f"Error generating detailed report: {str(e)}")
        
        file_stream = BytesIO()
        simple_report.save(file_stream)
        file_stream.seek(0)
        return file_stream